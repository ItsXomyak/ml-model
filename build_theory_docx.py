"""Генератор полной теоретической записки проекта в формате .docx.

Запуск:
    python build_theory_docx.py

На выходе: theory_full.docx — содержит развёрнутую теорию по всем темам,
покрытым в проекте: ML-основы, регрессия, кластеризация, регуляризация,
переобучение, метрики, MLOps, FastAPI, архитектура.
"""

import sys
from pathlib import Path

# UTF-8 stdout (важно для Windows-консоли)
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = Path(__file__).parent / "theory_full.docx"


def set_cell_bg(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_code(doc: Document, code: str) -> None:
    """Добавляет блок кода моноширинным шрифтом со светлым фоном."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    # Серый фон для блока кода через shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list, style: str = "List Bullet") -> None:
    for item in items:
        p = doc.add_paragraph(style=style)
        run = p.add_run(item)
        run.font.size = Pt(11)


def add_table(doc: Document, headers: list, rows: list, col_widths_cm=None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Заголовки
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        set_cell_bg(hdr[i], "D9E2F3")

    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for c, val in enumerate(row):
            cells[c].text = ""
            p = cells[c].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)

    if col_widths_cm:
        for col_idx, w in enumerate(col_widths_cm):
            for cell in table.columns[col_idx].cells:
                cell.width = Cm(w)


def build() -> None:
    doc = Document()

    # ===== Стиль документа =====
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ===== Титул =====
    title = doc.add_heading("Теоретическая записка к проекту", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("ML-система оценки стоимости и определения класса квартир")
    r.bold = True
    r.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Рубежный контроль · Дисциплина «Машинное обучение»\n"
        "Регрессия + кластеризация на синтетике рынка недвижимости г. Астаны\n"
        "Стек: Python 3.9+, scikit-learn, FastAPI, Docker"
    ).italic = True

    doc.add_page_break()

    # ===== Содержание (вручную) =====
    doc.add_heading("Содержание", level=1)
    toc = [
        "1. Введение и цели проекта",
        "2. Базовые понятия машинного обучения",
        "3. Структура датасета и предобработка",
        "4. Разделение данных: train / val / test (Задание 1)",
        "5. Регрессия: предсказание цены (Задание 3)",
        "6. Регуляризация: Lasso (L1) и Ridge (L2) (Задание 5)",
        "7. Переобучение и борьба с ним (Задание 6)",
        "8. Кластеризация k-means (Задание 4)",
        "9. Elbow method и Silhouette score",
        "10. Метрики качества моделей",
        "11. Воспроизводимость и random_state",
        "12. Сериализация моделей (joblib)",
        "13. MLOps: упаковка в Docker (Задание 7)",
        "14. Веб-сервис: FastAPI + статика",
        "15. Архитектура проекта",
        "16. Соответствие заданиям рубежного контроля",
        "17. Глоссарий",
    ]
    for item in toc:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item.split(". ", 1)[1]).font.size = Pt(11)

    doc.add_page_break()

    # ===== 1. Введение =====
    doc.add_heading("1. Введение и цели проекта", level=1)
    add_para(doc,
        "Проект решает две связанные задачи на синтетическом датасете рынка "
        "недвижимости г. Астаны: предсказание стоимости квартиры по её "
        "характеристикам (регрессия) и автоматическое определение класса "
        "квартиры — эконом / комфорт / бизнес / премиум — без ручной разметки "
        "методом k-means (кластеризация)."
    )
    add_para(doc, "Цели:", bold=True)
    add_bullets(doc, [
        "Закрыть все 8 заданий рубежного контроля (Блоки 1–4: split, регрессия, "
        "кластеризация, регуляризация, теория по overfitting, Dockerfile, инструкции).",
        "Получить воспроизводимую модель, обученную и упакованную одной командой docker run.",
        "Продемонстрировать корректную работу с переобучением через регуляризацию.",
        "Поднять веб-интерфейс для интерактивной демонстрации модели.",
    ])

    add_para(doc, "Стек технологий:", bold=True)
    add_bullets(doc, [
        "Python 3.9+ — язык реализации.",
        "pandas, numpy — работа с данными.",
        "scikit-learn — модели, препроцессинг, метрики.",
        "joblib — сериализация обученных моделей.",
        "matplotlib — визуализация графика elbow + silhouette.",
        "FastAPI + uvicorn — REST API и раздача фронта.",
        "Docker — упаковка системы для воспроизводимого запуска.",
    ])

    # ===== 2. Базовые понятия ML =====
    doc.add_heading("2. Базовые понятия машинного обучения", level=1)

    doc.add_heading("2.1. Что такое машинное обучение", level=2)
    add_para(doc,
        "Машинное обучение (ML) — раздел искусственного интеллекта, в котором "
        "алгоритм извлекает закономерности из данных и применяет их к новым "
        "наблюдениям. В отличие от классического программирования, где правила "
        "задаёт человек, в ML правила выводит сам алгоритм по примерам."
    )

    doc.add_heading("2.2. Виды задач", level=2)
    add_table(doc,
        ["Тип задачи", "Что предсказываем", "Пример", "В проекте"],
        [
            ["Регрессия (supervised)",       "Непрерывное число",        "Цена квартиры",        "✓ задача 1"],
            ["Классификация (supervised)",   "Метка из набора классов",  "Спам / не спам",       "—"],
            ["Кластеризация (unsupervised)", "Группировка без меток",    "Сегментация клиентов", "✓ задача 2"],
            ["Снижение размерности",         "Сжатие признаков",          "PCA, t-SNE",           "—"],
        ],
        col_widths_cm=[3.5, 4, 3.5, 2.5]
    )
    add_para(doc, "")
    add_para(doc,
        "В проекте используются обе парадигмы: регрессия (с учителем — у нас "
        "есть колонка price) и кластеризация (без учителя — мы НЕ говорим "
        "алгоритму, какая квартира к какому классу относится; он определяет "
        "это сам по структуре признаков)."
    )

    doc.add_heading("2.3. Жизненный цикл ML-задачи", level=2)
    add_bullets(doc, [
        "Постановка задачи и сбор данных.",
        "Разведочный анализ (EDA) — понимание распределений и связей.",
        "Препроцессинг — кодирование категорий, масштабирование, заполнение пропусков.",
        "Разделение на train / val / test.",
        "Обучение модели и подбор гиперпараметров на val.",
        "Финальная оценка на test (один раз!).",
        "Деплой: сериализация модели, упаковка в сервис.",
    ])

    # ===== 3. Структура датасета =====
    doc.add_heading("3. Структура датасета и предобработка", level=1)

    doc.add_heading("3.1. Признаки и целевая переменная", level=2)
    add_table(doc,
        ["Признак", "Тип", "Описание", "Диапазон"],
        [
            ["area",           "float",    "Площадь, м²",            "25 – 200"],
            ["rooms",          "int",      "Количество комнат",      "1 – 5"],
            ["floor",          "int",      "Этаж",                   "1 – 25"],
            ["total_floors",   "int",      "Этажей в доме",          "5 – 25"],
            ["year_built",     "int",      "Год постройки",          "1960 – 2025"],
            ["district",       "category", "Район Астаны",           "5 значений"],
            ["material",       "category", "Материал стен",          "панель/кирпич/монолит"],
            ["renovation",     "category", "Тип ремонта",            "4 значения"],
            ["dist_to_center", "float",    "До центра, км",          "0.5 – 25"],
            ["price",          "float",    "Целевая, млн ₸",         "15 – 250"],
        ],
        col_widths_cm=[3, 1.5, 5, 4]
    )

    doc.add_heading("3.2. Зачем синтетика", level=2)
    add_para(doc,
        "В проекте намеренно используется синтетический датасет, генерируемый "
        "скриптом generate_data.py. Преимущества:"
    )
    add_bullets(doc, [
        "Полный контроль над распределениями признаков и шумом.",
        "Гарантированная воспроизводимость (numpy.random.seed(42)).",
        "Известные «заложенные» зависимости — можно проверить, что модель их выучила.",
        "Внутри генерации заложены 4 латентных сегмента, которые кластеризатор должен обнаружить сам.",
    ])
    add_para(doc,
        "Цена в синтетике формируется как функция площади, базовой стоимости "
        "м² по сегменту, мультипликаторов ремонта/материала/года/расстояния/"
        "этажа и rooms_factor с гауссовским шумом ~5%."
    )

    doc.add_heading("3.3. Препроцессинг признаков", level=2)
    add_para(doc, "Числовые признаки", bold=True)
    add_para(doc,
        "Масштабируются через StandardScaler — приводятся к среднему 0 и "
        "стандартному отклонению 1: x_new = (x − μ) / σ. Зачем:"
    )
    add_bullets(doc, [
        "Lasso/Ridge чувствительны к масштабу: признак с большой амплитудой "
        "получит маленький коэффициент, что искажает регуляризацию.",
        "Алгоритмы на расстояниях (k-means, KNN) дают одинаковый вес всем измерениям.",
        "Градиентный спуск сходится быстрее на масштабированных данных.",
    ])

    add_para(doc, "Категориальные признаки", bold=True)
    add_para(doc,
        "Кодируются через OneHotEncoder — каждое значение становится "
        "отдельным бинарным столбцом. Например, district с 5 значениями → "
        "5 колонок district_Есиль, district_Алматы, ..."
    )
    add_para(doc, "Альтернативы:", bold=True)
    add_bullets(doc, [
        "Label encoding — присваивает каждому значению целое число (0,1,2,...). "
        "Подходит, когда категория ОРДИНАЛЬНА (упорядочена). В проекте используется "
        "для district при кластеризации — ранжируем по средней цене м² внутри района.",
        "Target encoding — заменяем категорию средним значением target по этой категории. "
        "Мощно, но риск утечки.",
        "Embeddings — для глубоких моделей; превращает категорию в вектор низкой размерности.",
    ])

    add_para(doc, "ColumnTransformer + Pipeline", bold=True)
    add_para(doc,
        "Чтобы применять разные преобразования к разным колонкам, используется "
        "ColumnTransformer. Чтобы объединить препроцессинг с моделью в один "
        "объект — Pipeline. Это даёт два преимущества:"
    )
    add_bullets(doc, [
        "Препроцессинг применяется одинаково к train, val, test, на проде — "
        "невозможно «забыть» масштабировать вход.",
        "Один вызов pipeline.fit_transform(X_train) обучает и StandardScaler, "
        "и OneHotEncoder, и модель — атомарно.",
    ])

    add_code(doc,
        "preprocessor = ColumnTransformer([\n"
        "    ('num', StandardScaler(), NUMERIC_FEATURES),\n"
        "    ('cat', OneHotEncoder(handle_unknown='ignore'), CATEGORICAL_FEATURES),\n"
        "])\n"
        "pipeline = Pipeline([\n"
        "    ('preproc', preprocessor),\n"
        "    ('model',   RandomForestRegressor(max_depth=15)),\n"
        "])"
    )

    # ===== 4. Train/Val/Test =====
    doc.add_heading("4. Разделение данных: train / val / test (Задание 1)", level=1)

    doc.add_heading("4.1. Назначение трёх выборок", level=2)
    add_table(doc,
        ["Выборка", "Доля", "Используется для", "Видна модели"],
        [
            ["train", "60%", "Обучения параметров (весов)",                              "Многократно"],
            ["val",   "20%", "Подбора гиперпараметров (alpha, max_depth, K)",            "Многократно, без обучения"],
            ["test",  "20%", "Финальной оценки качества после фиксации модели",          "ОДИН раз в конце"],
        ],
        col_widths_cm=[2.5, 1.5, 7, 4]
    )

    doc.add_heading("4.2. В чём разница между val и test", level=2)
    add_bullets(doc, [
        "train — здесь алгоритм НАСТРАИВАЕТ свои параметры (веса, деревья и т.д.).",
        "val — здесь МЫ принимаем решения: какой alpha взять, какая глубина "
        "леса, сколько кластеров. Алгоритм НЕ учится напрямую на val, но мы "
        "выбираем лучшую конфигурацию по метрике на val.",
        "test — НЕТРОНУТЫЙ контрольный набор. Используется один раз в самом "
        "конце, после того как все решения приняты.",
    ])

    doc.add_heading("4.3. Почему нельзя тюнить гиперпараметры по test", level=2)
    add_para(doc,
        "Если подбирать гиперпараметры по test, возникает четыре проблемы:"
    )
    add_bullets(doc, [
        "Утечка информации: каждый раз глядя на качество на test и меняя модель, "
        "мы неявно подгоняем модель под этот test. Test перестаёт быть независимым.",
        "Иллюзия высокого качества: среди множества переборов мы найдём ту, "
        "которая случайно лучше всего легла на этот конкретный test. Метрика будет завышена.",
        "Невозможность честной оценки: test становится «знакомым» модели, как val. "
        "Мы теряем единственный беспристрастный набор данных.",
        "Аналогия с экзаменом: train — лекции, val — контрольные, test — экзамен. "
        "Если студент видит экзаменационные билеты заранее — оценка не отражает знания.",
    ])

    doc.add_heading("4.4. Двухступенчатый split", level=2)
    add_para(doc,
        "В коде train.py split реализован в два шага: сначала отделяем 60% "
        "train от 40% temp, затем temp делим пополам на val и test. "
        "random_state=42 фиксирует одно и то же разбиение при каждом запуске."
    )
    add_code(doc,
        "X_train, X_temp, y_train, y_temp = train_test_split(\n"
        "    X, y, test_size=0.40, random_state=42)\n"
        "X_val, X_test, y_val, y_test = train_test_split(\n"
        "    X_temp, y_temp, test_size=0.50, random_state=42)"
    )

    doc.add_heading("4.5. Стратификация", level=2)
    add_para(doc,
        "Для классификации используется stratify=y — разделение с сохранением "
        "пропорций классов в train/val/test. Для регрессии можно стратифицировать "
        "по биннингу таргета (pd.qcut), особенно когда редкие сегменты (премиум-"
        "квартиры в нашем датасете ~17%) могут случайно попасть только в train или test."
    )

    # ===== 5. Регрессия =====
    doc.add_heading("5. Регрессия: предсказание цены (Задание 3)", level=1)

    doc.add_heading("5.1. Постановка", level=2)
    add_para(doc,
        "Дано: 9 признаков квартиры. Найти: оценку цены в млн ₸. "
        "Это задача регрессии — предсказание непрерывной величины."
    )

    doc.add_heading("5.2. Линейная регрессия", level=2)
    add_para(doc,
        "Модель ищет линейную комбинацию признаков, минимизирующую среднеквадратичную ошибку:"
    )
    add_code(doc, "ŷ = w₀ + w₁·x₁ + w₂·x₂ + ... + wₙ·xₙ")
    add_para(doc, "Цель — найти веса w*, минимизирующие L(w) = Σ(yᵢ − ŷᵢ)².")
    add_para(doc, "Плюсы и минусы:", bold=True)
    add_bullets(doc, [
        "+ Простая и быстрая. Легко интерпретируется (w покажет важность признаков).",
        "+ Имеет аналитическое решение (метод наименьших квадратов).",
        "− Не моделирует нелинейные зависимости.",
        "− Чувствительна к выбросам.",
    ])

    doc.add_heading("5.3. RandomForestRegressor", level=2)
    add_para(doc,
        "Ансамбль решающих деревьев. Каждое дерево обучается на бутстрап-выборке "
        "и случайном подмножестве признаков. Финальный прогноз — среднее по всем деревьям."
    )
    add_para(doc, "Ключевые гиперпараметры:", bold=True)
    add_bullets(doc, [
        "n_estimators — число деревьев (больше = стабильнее, но дороже).",
        "max_depth — максимальная глубина дерева (ограничивает переобучение).",
        "min_samples_leaf — минимум объектов в листе.",
        "max_features — сколько признаков рассматривать в каждом сплите.",
    ])
    add_para(doc, "Плюсы и минусы:", bold=True)
    add_bullets(doc, [
        "+ Захватывает нелинейные зависимости и взаимодействия признаков.",
        "+ Робастен к выбросам и шуму.",
        "+ Не требует масштабирования (хотя в Pipeline это всё равно есть для совместимости).",
        "− Медленнее линейных моделей.",
        "− Хуже экстраполирует за пределы train (деревья дают константу в листе).",
    ])

    doc.add_heading("5.4. Что получилось в проекте", level=2)
    add_para(doc,
        "В train.py обучаются три кандидата в base-модели: LinearRegression и "
        "RandomForestRegressor с двумя глубинами (10 и 15). Лучшая по MSE на val — "
        "RandomForest(max_depth=15). На test финальные метрики:"
    )
    add_bullets(doc, [
        "R² = 0.9945 (требование ≥ 0.85 выполнено с большим запасом).",
        "MSE = 38.93 (RMSE ≈ 6.24 млн ₸ при средней цене ~107 млн).",
    ])

    # ===== 6. Регуляризация =====
    doc.add_heading("6. Регуляризация: Lasso (L1) и Ridge (L2) (Задание 5)", level=1)

    doc.add_heading("6.1. Идея регуляризации", level=2)
    add_para(doc,
        "Регуляризация — добавление штрафного слагаемого к функции потерь, "
        "чтобы ограничить рост весов и тем самым уменьшить переобучение. "
        "Без регуляризации линейная модель может «выкрутить» большие веса, "
        "идеально ложась на train, но плохо обобщая."
    )
    add_code(doc,
        "L(w) = Σ(yᵢ − ŷᵢ)² + α · R(w)\n"
        "α — коэффициент силы регуляризации\n"
        "R(w) — штрафное слагаемое"
    )

    doc.add_heading("6.2. L1 — Lasso", level=2)
    add_code(doc, "R(w) = Σ |wⱼ|")
    add_para(doc, "Свойства:", bold=True)
    add_bullets(doc, [
        "Производит разреженные решения: часть коэффициентов становится строго нулевыми.",
        "Фактически выполняет отбор признаков (feature selection).",
        "Геометрия штрафа — ромб (l1-шар), его углы лежат на осях, "
        "поэтому решение часто «прилипает» к нулю по некоторым координатам.",
        "Чувствителен к коллинеарности: из группы скоррелированных признаков выбирает один произвольно.",
    ])

    doc.add_heading("6.3. L2 — Ridge", level=2)
    add_code(doc, "R(w) = Σ wⱼ²")
    add_para(doc, "Свойства:", bold=True)
    add_bullets(doc, [
        "Уменьшает все веса, но НЕ зануляет их.",
        "Хорошо справляется с мультиколлинеарностью — распределяет вес между скоррелированными признаками.",
        "Геометрия штрафа — круг (l2-шар), без угловых точек, поэтому веса плавно сжимаются.",
        "Имеет аналитическое решение: w = (XᵀX + αI)⁻¹Xᵀy.",
    ])

    doc.add_heading("6.4. Сравнение L1 и L2", level=2)
    add_table(doc,
        ["Свойство", "L1 (Lasso)", "L2 (Ridge)"],
        [
            ["Штраф",                       "Σ|wⱼ|",                              "Σwⱼ²"],
            ["Зануляет веса",               "Да",                                 "Нет"],
            ["Feature selection",           "Да, автоматически",                  "Нет"],
            ["Мультиколлинеарность",        "Плохо обрабатывает",                 "Хорошо"],
            ["Аналитическое решение",       "Нет (coordinate descent)",            "Да"],
            ["Стабильность к шуму",         "Ниже",                               "Выше"],
            ["Интерпретируемость",          "Высокая (мало признаков)",            "Средняя"],
        ],
        col_widths_cm=[5, 4.5, 4.5]
    )

    doc.add_heading("6.5. ElasticNet — компромисс", level=2)
    add_code(doc, "R(w) = ρ · Σ|wⱼ| + (1 − ρ) · Σwⱼ²")
    add_para(doc,
        "Линейная комбинация L1 и L2. Используется, когда есть и желание "
        "зануления признаков, и проблема коллинеарности."
    )

    doc.add_heading("6.6. Когда использовать Lasso", level=2)
    add_bullets(doc, [
        "Много признаков, часть из которых заведомо неинформативна.",
        "Нужна интерпретируемая модель (мало ненулевых коэффициентов).",
        "Высокоразмерные задачи (p ≫ n) — генетика, тексты.",
        "Гипотеза о разреженности истинной модели — целевую переменную "
        "определяют 2-3 ключевых фактора, а остальное шум.",
    ])

    doc.add_heading("6.7. Когда использовать Ridge", level=2)
    add_bullets(doc, [
        "Все признаки потенциально полезны, нужно лишь стабилизировать модель.",
        "Сильная мультиколлинеарность (например, area и rooms сильно коррелируют).",
        "Нужно гладкое аналитическое поведение.",
    ])

    doc.add_heading("6.8. Что в проекте", level=2)
    add_para(doc,
        "В train_regularized() перебирается 7 значений alpha [0.001, 0.01, "
        "0.1, 0.5, 1, 5, 10] для Lasso и Ridge. Лучшая модель выбирается "
        "по min MSE на val."
    )

    # ===== 7. Переобучение =====
    doc.add_heading("7. Переобучение и борьба с ним (Задание 6)", level=1)

    doc.add_heading("7.1. Что такое overfitting", level=2)
    add_para(doc,
        "Переобучение — состояние, когда модель запомнила train, включая шум, "
        "но плохо обобщает на новые данные. Симптом: высокая метрика на train "
        "и существенно более низкая на val/test."
    )
    add_para(doc, "Сценарий из задания: R² = 0.98 на train и R² = 0.72 на val. Разрыв 26 п.п. — классический overfit.")

    doc.add_heading("7.2. Bias / Variance trade-off", level=2)
    add_bullets(doc, [
        "Bias — систематическая ошибка из-за слишком простой модели (underfit).",
        "Variance — разброс предсказаний из-за чрезмерной чувствительности к train (overfit).",
        "Цель — найти баланс: и bias, и variance — низкие.",
    ])

    doc.add_heading("7.3. Три шага исправления overfit", level=2)
    add_para(doc, "Шаг 1. Усилить регуляризацию", bold=True)
    add_bullets(doc, [
        "Lasso/Ridge: увеличить alpha (например, с 0.1 до 1.0 или 10).",
        "RandomForest/GradientBoosting: уменьшить max_depth, увеличить min_samples_leaf, уменьшить n_estimators.",
        "Эффект: train-метрика немного просядет, val вырастет, разрыв сократится.",
    ])

    add_para(doc, "Шаг 2. Уменьшить сложность признакового пространства", bold=True)
    add_bullets(doc, [
        "Удалить шумовые признаки — Lasso уже сигналит, какие занулять.",
        "Убрать сильно коллинеарные пары (одну из них).",
        "Снизить размерность через PCA (95% дисперсии).",
        "Не использовать высокие степени полиномов или сложные feature engineering без необходимости.",
    ])

    add_para(doc, "Шаг 3. Увеличить объём данных или применить cross-validation", bold=True)
    add_bullets(doc, [
        "Собрать/сгенерировать больше данных (--n с 3000 до 10000).",
        "Аугментация: реалистично шумить признаки.",
        "K-fold cross-validation: вместо одной val — 5 или 10 фолдов.",
        "Early stopping для итеративных алгоритмов.",
        "Данные — самый сильный регуляризатор.",
    ])

    doc.add_heading("7.4. Целевое состояние после фиксов", level=2)
    add_bullets(doc, [
        "Разрыв train − val ≤ 5–10 процентных пунктов.",
        "val-метрика приемлемо высока (для ТЗ — R² ≥ 0.85).",
        "Если после трёх шагов всё ещё overfit — выбран слишком гибкий класс моделей; перейти на проще.",
    ])

    # ===== 8. Кластеризация =====
    doc.add_heading("8. Кластеризация k-means (Задание 4)", level=1)

    doc.add_heading("8.1. Постановка", level=2)
    add_para(doc,
        "Кластеризация — обучение БЕЗ учителя. У нас нет колонки «класс квартиры». "
        "Алгоритм должен сам найти группы похожих объектов."
    )

    doc.add_heading("8.2. Алгоритм k-means", level=2)
    add_bullets(doc, [
        "Шаг 1. Выбираем число кластеров K и инициализируем K случайных центроидов.",
        "Шаг 2. Каждой точке присваиваем номер ближайшего центроида (E-step).",
        "Шаг 3. Пересчитываем центроиды как среднее точек кластера (M-step).",
        "Шаг 4. Повторяем 2–3, пока центроиды не стабилизируются.",
    ])

    doc.add_heading("8.3. Inertia (WCSS)", level=2)
    add_code(doc, "Inertia = Σᵢ ||xᵢ − μ_{c(i)}||²")
    add_para(doc,
        "Сумма квадратов расстояний от каждой точки до центроида её кластера. "
        "Чем меньше K, тем больше inertia. При K = N (число точек) inertia = 0. "
        "Поэтому абсолютное значение бессмысленно — важен темп убывания."
    )

    doc.add_heading("8.4. Признаки для k-means", level=2)
    add_para(doc,
        "В проекте используются: цена за м² (price/area), площадь, расстояние "
        "до центра, ранг района (label-encoded по средней цене м² в районе). "
        "Все стандартизируются, потому что k-means на евклидовом расстоянии."
    )

    doc.add_heading("8.5. Маппинг кластеров на классы", level=2)
    add_para(doc, "После обучения с K=4:")
    add_bullets(doc, [
        "Считаем среднюю price_per_m2 в каждом кластере.",
        "Сортируем кластеры от дешёвого к дорогому.",
        "Присваиваем имена: эконом → комфорт → бизнес → премиум.",
        "Если K ≠ 4 — используем generic «класс 1, ..., класс N».",
    ])

    # ===== 9. Elbow + silhouette =====
    doc.add_heading("9. Elbow method и Silhouette score", level=1)

    doc.add_heading("9.1. Метод локтя", level=2)
    add_para(doc,
        "Перебираем K от 1 до 10, считаем inertia, ищем «излом» (точка, где "
        "темп убывания резко падает). Реализован геометрический подход:"
    )
    add_bullets(doc, [
        "Считаем drops[k] = inertia(K=k) − inertia(K=k+1).",
        "Считаем ratio[k] = drops[k+1] / drops[k] — насколько следующее улучшение меньше предыдущего.",
        "Первое отношение пропускаем (доминирующее падение K=1→K=2 всегда искажает).",
        "Минимум среди оставшихся = точка максимального замедления = optimal K.",
    ])
    add_para(doc,
        "Этот алгоритм устойчивее «перпендикулярного расстояния» от линии "
        "соединяющей крайние точки, который смещается к ранним K из-за резкого "
        "первого drop."
    )

    doc.add_heading("9.2. Silhouette score", level=2)
    add_code(doc,
        "s(i) = (b(i) − a(i)) / max(a(i), b(i))\n"
        "a(i) — среднее расстояние от точки до её кластера\n"
        "b(i) — среднее расстояние до ближайшего ЧУЖОГО кластера"
    )
    add_para(doc, "Интерпретация:")
    add_bullets(doc, [
        "+1 — кластеры компактны и хорошо разделены.",
        "0 — границы кластеров размыты.",
        "−1 — точка ошибочно отнесена к чужому кластеру.",
    ])
    add_para(doc,
        "Используется как вторая метрика. Если elbow и пик silhouette сходятся "
        "на одном K — выбор подтверждён. В нашем проекте при K=4 silhouette ≈ "
        "0.59 (максимум среди K≥2)."
    )

    # ===== 10. Метрики =====
    doc.add_heading("10. Метрики качества моделей", level=1)

    doc.add_heading("10.1. Метрики регрессии", level=2)
    add_table(doc,
        ["Метрика", "Формула", "Интерпретация"],
        [
            ["MSE",  "(1/n)Σ(yᵢ − ŷᵢ)²",                       "Средний квадрат ошибки. Штрафует большие ошибки."],
            ["RMSE", "√MSE",                                    "Корень из MSE. В тех же единицах, что target."],
            ["MAE",  "(1/n)Σ|yᵢ − ŷᵢ|",                         "Средняя абсолютная ошибка. Робастнее к выбросам."],
            ["R²",   "1 − Σ(yᵢ−ŷᵢ)² / Σ(yᵢ−ȳ)²",                "Доля объяснённой дисперсии. 1 — идеал, 0 — модель не лучше среднего."],
        ],
        col_widths_cm=[2, 5, 7]
    )
    add_para(doc, "В проекте обязательно по ТЗ — MSE и R².")

    doc.add_heading("10.2. Метрики кластеризации", level=2)
    add_bullets(doc, [
        "Inertia — внутрикластерная компактность (используется для elbow).",
        "Silhouette — внешняя оценка качества разбиения.",
        "Davies-Bouldin index — соотношение внутрикластерной и межкластерной дисперсии.",
        "Calinski-Harabasz — отношение межкластерной дисперсии к внутрикластерной.",
    ])

    # ===== 11. random_state =====
    doc.add_heading("11. Воспроизводимость и random_state", level=1)
    add_para(doc,
        "Воспроизводимость — гарантия, что повторный запуск даёт идентичный "
        "результат. Достигается фиксацией всех источников случайности через "
        "seed (random_state)."
    )
    add_para(doc, "Где в проекте важен random_state=42:")
    add_bullets(doc, [
        "numpy.random.default_rng(42) в generate_data.py — сама синтетика.",
        "train_test_split(random_state=42) — одинаковое разбиение train/val/test.",
        "RandomForestRegressor(random_state=42) — детерминированные бутстрап-выборки и сплиты.",
        "Lasso/Ridge(random_state=42) — coordinate descent инициализация.",
        "KMeans(random_state=42) — фиксированная инициализация центроидов.",
    ])
    add_para(doc,
        "Требование ТЗ: повторный запуск python train.py даёт идентичные "
        "метрики до 4-го знака. Это критично для производственных ML-систем — "
        "без воспроизводимости невозможно отладить регрессию качества."
    )

    # ===== 12. joblib =====
    doc.add_heading("12. Сериализация моделей (joblib)", level=1)
    add_para(doc,
        "После обучения объект Pipeline нужно сохранить на диск, чтобы потом "
        "загружать без переобучения. В sklearn для этого используется joblib."
    )
    add_code(doc,
        "import joblib\n"
        "joblib.dump(pipeline, 'models/price_model.pkl')\n"
        "model = joblib.load('models/price_model.pkl')\n"
        "predictions = model.predict(X_new)"
    )
    add_para(doc, "joblib vs pickle:", bold=True)
    add_bullets(doc, [
        "joblib эффективнее на numpy-массивах (используется в обученных моделях).",
        "joblib умеет хранить большие объекты в отдельных файлах.",
        "Pickle универсальнее, но медленнее на массивах.",
    ])
    add_para(doc, "Что сохраняется:", bold=True)
    add_bullets(doc, [
        "Все шаги Pipeline (StandardScaler, OneHotEncoder, модель).",
        "Параметры обученных трансформеров (μ, σ, словари кодирования).",
        "Параметры модели (веса, деревья).",
    ])
    add_para(doc,
        "Важно: если вы обучили StandardScaler на train, при инференсе он "
        "должен применять ТЕ ЖЕ μ и σ. Pipeline + joblib гарантируют это автоматически."
    )

    # ===== 13. Docker =====
    doc.add_heading("13. MLOps: упаковка в Docker (Задание 7)", level=1)

    doc.add_heading("13.1. Зачем Docker для ML", level=2)
    add_bullets(doc, [
        "Воспроизводимость окружения: «у меня работает» больше не отговорка.",
        "Изоляция от хост-системы: ваш Python 3.9 не конфликтует с системным Python.",
        "Простой деплой: docker run на любом сервере с Docker.",
        "CI/CD совместимость: образ — артефакт сборки.",
    ])

    doc.add_heading("13.2. Dockerfile проекта", level=2)
    add_code(doc,
        "FROM python:3.9-slim\n"
        "WORKDIR /app\n"
        "COPY requirements.txt .\n"
        "RUN pip install --no-cache-dir -r requirements.txt\n"
        "COPY . .\n"
        "CMD [\"python\", \"train.py\"]"
    )
    add_para(doc, "Разбор инструкций:", bold=True)
    add_bullets(doc, [
        "FROM python:3.9-slim — базовый образ. slim экономит ~700 МБ vs full image.",
        "WORKDIR /app — рабочая директория внутри контейнера.",
        "COPY requirements.txt . — копирование сначала только requirements (для кеша слоёв).",
        "RUN pip install --no-cache-dir — установка зависимостей. --no-cache-dir экономит размер образа.",
        "COPY . . — копирование остального кода (после deps, чтобы изменения кода не инвалидировали кеш слоя deps).",
        "CMD — команда по умолчанию при docker run.",
    ])

    doc.add_heading("13.3. Команды сборки и запуска (Задание 8)", level=2)
    add_code(doc,
        "docker build -t ml-model-v1 .\n"
        "docker run --rm ml-model-v1"
    )
    add_para(doc,
        "build собирает образ из Dockerfile. -t задаёт тег. run запускает "
        "контейнер. --rm удаляет контейнер после завершения (артефакты "
        "остаются в образе)."
    )

    doc.add_heading("13.4. .dockerignore", level=2)
    add_para(doc,
        "Аналог .gitignore для Docker. Исключает из контекста сборки лишние "
        "файлы (.git, *.pyc, виртуальные окружения), уменьшая размер образа "
        "и ускоряя сборку."
    )

    # ===== 14. FastAPI =====
    doc.add_heading("14. Веб-сервис: FastAPI + статика", level=1)

    doc.add_heading("14.1. Архитектура", level=2)
    add_para(doc,
        "После обучения модели поднимается REST-сервис. Один Python-процесс "
        "обслуживает API-эндпоинты и раздаёт фронт. Внутри:"
    )
    add_bullets(doc, [
        "FastAPI — современный async веб-фреймворк со встроенной валидацией через Pydantic.",
        "uvicorn — ASGI-сервер, на котором FastAPI запускается.",
        "Pydantic — модели данных с автоматической валидацией типов и диапазонов.",
        "StaticFiles — раздача CSS/JS/картинок.",
        "FileResponse — отдача index.html на корне.",
    ])

    doc.add_heading("14.2. Endpoints", level=2)
    add_table(doc,
        ["Метод", "Путь", "Описание"],
        [
            ["GET",  "/",            "Веб-форма (index.html)"],
            ["GET",  "/static/*",    "CSS / JS / изображения"],
            ["GET",  "/api/health",  "Статус сервера, метрики модели, список районов"],
            ["POST", "/api/predict", "Цена и класс квартиры по JSON"],
            ["GET",  "/docs",        "Автогенерированный Swagger UI"],
        ],
        col_widths_cm=[1.5, 3, 9]
    )

    doc.add_heading("14.3. Pydantic-валидация", level=2)
    add_para(doc,
        "Каждый запрос на /api/predict валидируется автоматически — диапазоны "
        "признаков (area 25-200, rooms 1-5 и т.д.) проверяются на стороне сервера. "
        "При нарушении возвращается 422 Unprocessable Entity с понятным "
        "сообщением о том, какое поле не прошло валидацию."
    )

    doc.add_heading("14.4. Lazy-loading модели", level=2)
    add_para(doc,
        "Модели загружаются один раз при первом запросе (load_artifacts). "
        "Это ускоряет старт сервера и позволяет менять модели без рестарта."
    )

    # ===== 15. Архитектура =====
    doc.add_heading("15. Архитектура проекта", level=1)
    add_code(doc,
        "ml-model/\n"
        "├── data/\n"
        "│   └── apartments_astana.csv      ← синтетика (генерируется)\n"
        "├── models/                          ← создаётся при обучении\n"
        "│   ├── price_model.pkl              ← regressor pipeline\n"
        "│   └── cluster_model.pkl            ← KMeans + StandardScaler\n"
        "├── static/                          ← фронт\n"
        "│   ├── index.html\n"
        "│   ├── style.css\n"
        "│   └── script.js\n"
        "├── generate_data.py                 ← генератор синтетики\n"
        "├── train.py                         ← главный обучающий скрипт\n"
        "├── app.py                           ← FastAPI: API + статика\n"
        "├── theory.md / theory_full.docx     ← теория\n"
        "├── plots_explanation.md             ← пояснение к elbow_plot.png\n"
        "├── elbow_plot.png                   ← визуализация elbow + silhouette\n"
        "├── requirements.txt                 ← Python-зависимости\n"
        "├── Dockerfile                       ← упаковка системы\n"
        "├── README.md                        ← инструкции\n"
        "└── metrics.json                     ← все метрики обученной модели"
    )

    add_para(doc, "Поток данных:", bold=True)
    add_bullets(doc, [
        "1. generate_data.py → data/apartments_astana.csv",
        "2. train.py → читает CSV → split → обучение → models/*.pkl + metrics.json + elbow_plot.png",
        "3. app.py → joblib.load(models/*.pkl) → принимает JSON → возвращает цену + класс",
        "4. static/index.html → форма → fetch /api/predict → отрисовка результата",
    ])

    # ===== 16. Соответствие заданиям =====
    doc.add_heading("16. Соответствие заданиям рубежного контроля", level=1)
    add_table(doc,
        ["№", "Описание", "Где реализовано"],
        [
            ["1", "Теория: val vs test",         "theory.md / раздел 4"],
            ["2", "Split 60/20/20",               "train.py: split_data()"],
            ["3", "Регрессия + MSE и R²",         "train.py: train_regression()"],
            ["4", "k-means + elbow method",       "train.py: train_clustering()"],
            ["5", "Lasso vs Ridge + код",         "train.py: train_regularized() + раздел 6"],
            ["6", "3 шага при overfit",           "theory.md / раздел 7"],
            ["7", "Dockerfile",                   "Dockerfile"],
            ["8", "Команды build/run",            "README.md"],
        ],
        col_widths_cm=[1, 6, 8]
    )

    # ===== 17. Глоссарий =====
    doc.add_heading("17. Глоссарий", level=1)
    glossary = [
        ("Feature (признак)",          "Колонка во входных данных модели."),
        ("Target (целевая)",           "Колонка, которую модель предсказывает."),
        ("Hyperparameter",             "Параметр алгоритма, задаваемый ДО обучения (alpha, max_depth, K)."),
        ("Parameter (вес)",            "Параметр модели, ОБУЧАЕМЫЙ алгоритмом (веса w в линрегрессии)."),
        ("Pipeline",                    "Последовательность шагов препроцессинга и модели как один объект."),
        ("ColumnTransformer",          "Контейнер sklearn, применяющий разные трансформеры к разным колонкам."),
        ("StandardScaler",             "Z-нормализация: x_new = (x − μ) / σ."),
        ("OneHotEncoder",              "Кодирование категории в N бинарных колонок."),
        ("Label Encoding",             "Кодирование категории целым числом."),
        ("MSE",                        "Mean Squared Error — средний квадрат ошибки."),
        ("R² (коэф. детерминации)",     "Доля объяснённой дисперсии. 1 — идеал, 0 — на уровне среднего."),
        ("Inertia / WCSS",             "Within-Cluster Sum of Squares — сумма квадратов расстояний до центроида."),
        ("Silhouette",                 "Метрика качества кластеризации в [-1, 1]."),
        ("Elbow method",               "Эвристика выбора K через излом кривой inertia."),
        ("Overfitting",                "Переобучение: высокая метрика на train, низкая на val/test."),
        ("Underfitting",               "Недообучение: модель слишком простая для задачи."),
        ("Bias",                       "Систематическая ошибка модели."),
        ("Variance",                   "Чувствительность к train (разброс)."),
        ("Regularization",             "Штраф за сложность для уменьшения overfit."),
        ("L1 (Lasso)",                 "Σ|w| — даёт разреженность."),
        ("L2 (Ridge)",                 "Σw² — сжимает все веса."),
        ("ElasticNet",                 "Комбинация L1 и L2."),
        ("Cross-validation (CV)",      "Множественная валидация на K-фолдах."),
        ("random_state / seed",        "Зерно генератора случайных чисел для воспроизводимости."),
        ("joblib",                     "Библиотека сериализации numpy-моделей."),
        ("Docker",                     "Контейнеризация — упаковка приложения с окружением."),
        ("Dockerfile",                 "Декларативное описание сборки Docker-образа."),
        ("Image vs Container",         "Image — шаблон, Container — запущенный экземпляр."),
        ("FastAPI",                    "Современный async-фреймворк REST-API на Python."),
        ("ASGI (uvicorn)",             "Async Server Gateway Interface — спецификация серверов для async-приложений."),
        ("Pydantic",                   "Библиотека валидации данных через type hints."),
        ("REST API",                   "HTTP-интерфейс с CRUD-операциями над ресурсами."),
        ("CRUD",                       "Create, Read, Update, Delete — четыре базовые операции."),
        ("MLOps",                      "Практики DevOps для ML: воспроизводимость, мониторинг, деплой моделей."),
    ]
    for term, definition in glossary:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        run_term = p.add_run(term + " — ")
        run_term.bold = True
        run_term.font.size = Pt(11)
        run_def = p.add_run(definition)
        run_def.font.size = Pt(11)

    # ===== Конец =====
    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = end.add_run("— Конец теоретической записки —")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.save(OUT)
    print(f"[build_theory_docx] Сохранено → {OUT}")
    print(f"[build_theory_docx] Размер: {OUT.stat().st_size / 1024:.1f} КБ")


if __name__ == "__main__":
    build()
